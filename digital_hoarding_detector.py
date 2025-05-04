import os
import zipfile
import rarfile
import py7zr
import hashlib
import difflib
from docx import Document
import PyPDF2
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from io import BytesIO
from PyPDF2 import PdfReader

DOWNLOADS_PATH = os.path.expanduser("~/Downloads")
SPECIAL_EXTENSIONS = ['.zip', '.rar', '.7z']

NORMAL_FILES, NORMAL_FOLDERS, ARCHIVE_SUMMARY = [], [], []
TOTAL_FILES_IN_ARCHIVES = 0
TOTAL_FOLDERS_IN_ARCHIVES = 0
MAX_NESTING_DEPTH = 0

def scan_normal_filesystem(target_folder):
    for root, dirs, files in os.walk(target_folder):
        for dir in dirs:
            NORMAL_FOLDERS.append(os.path.join(root, dir))
        for file in files:
            file_path = os.path.join(root, file)
            ext = os.path.splitext(file)[1].lower()
            if ext in SPECIAL_EXTENSIONS:
                scan_archive(file_path, 1)
            else:
                NORMAL_FILES.append(file_path)

def scan_archive(archive_path, current_level):
    global TOTAL_FILES_IN_ARCHIVES, TOTAL_FOLDERS_IN_ARCHIVES, MAX_NESTING_DEPTH
    try:
        if archive_path.endswith('.zip'):
            with zipfile.ZipFile(archive_path, 'r') as archive:
                for info in archive.infolist():
                    if info.is_dir():
                        TOTAL_FOLDERS_IN_ARCHIVES += 1
                    else:
                        TOTAL_FILES_IN_ARCHIVES += 1
                    nesting = info.filename.count('/')
                    MAX_NESTING_DEPTH = max(MAX_NESTING_DEPTH, current_level + nesting)
                ARCHIVE_SUMMARY.append(archive_path)
        elif archive_path.endswith('.rar'):
            with rarfile.RarFile(archive_path) as archive:
                for info in archive.infolist():
                    if info.isdir():
                        TOTAL_FOLDERS_IN_ARCHIVES += 1
                    else:
                        TOTAL_FILES_IN_ARCHIVES += 1
                    nesting = info.filename.count('/')
                    MAX_NESTING_DEPTH = max(MAX_NESTING_DEPTH, current_level + nesting)
                ARCHIVE_SUMMARY.append(archive_path)
        elif archive_path.endswith('.7z'):
            with py7zr.SevenZipFile(archive_path, mode='r') as archive:
                for info in archive.list():
                    if info.is_directory:
                        TOTAL_FOLDERS_IN_ARCHIVES += 1
                    else:
                        TOTAL_FILES_IN_ARCHIVES += 1
                    nesting = info.filename.count('/')
                    MAX_NESTING_DEPTH = max(MAX_NESTING_DEPTH, current_level + nesting)
                ARCHIVE_SUMMARY.append(archive_path)
    except Exception as e:
        print(f"⚠️ Skipping archive (could not open): {archive_path}")

def calculate_hash(filepath):
    hash_sha256 = hashlib.sha256()
    try:
        with open(filepath, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_sha256.update(chunk)
    except Exception as e:
        print(f"⚠️ Could not read {filepath}: {e}")
    return hash_sha256.hexdigest()

def ask_file_deletion(file1, file2):
    print("\n🗑️ What do you want to do with these duplicate files?")
    print("[1] Keep both")
    print("[2] Delete first file")
    print("[3] Delete second file")
    print("[4] Delete both")
    while True:
        choice = input("Your choice [1-4]: ").strip()
        if choice == '1':
            print("✅ Keeping both files.")
            break
        elif choice == '2':
            os.remove(file1)
            print(f"🗑️ Deleted: {file1}")
            break
        elif choice == '3':
            os.remove(file2)
            print(f"🗑️ Deleted: {file2}")
            break
        elif choice == '4':
            os.remove(file1)
            os.remove(file2)
            print("🗑️ Deleted both files.")
            break
        else:
            print("❌ Invalid choice. Enter 1, 2, 3, or 4.")

def compare_and_merge(file1, file2):
    try:
        with open(file1, 'r', encoding="utf-8", errors="ignore") as f1, open(file2, 'r', encoding="utf-8", errors="ignore") as f2:
            text1 = f1.read()
            text2 = f2.read()
    except Exception as e:
        print(f"⚠️ Could not read files:\n - {file1}\n - {file2}\n⚠️ Reason: {e}")
        return False

    if text1 == text2:
        print("🔁 Files are completely identical. No merge needed.")
        ask_file_deletion(file1, file2)
        return False

    print("🔍 Files differ. Showing character-level changes:")
    matcher = difflib.SequenceMatcher(None, text1, text2)
    result = []
    for opcode, i1, i2, j1, j2 in matcher.get_opcodes():
        a = text1[i1:i2]
        b = text2[j1:j2]
        if opcode == 'equal':
            result.append(a)
        elif opcode == 'insert':
            print(f"+ Insert: '{b}'")
            result.append(b)
        elif opcode == 'delete':
            print(f"- Delete: '{a}'")
        elif opcode == 'replace':
            print(f"~ Replace '{a}' with '{b}'")
            choice = input("Choose (1) keep A, (2) keep B, (3) both, (4) custom edit: ")
            if choice == '1':
                result.append(a)
            elif choice == '2':
                result.append(b)
            elif choice == '3':
                result.append(a + b)
            elif choice == '4':
                custom = input("Enter custom replacement: ")
                result.append(custom)
    merged = ''.join(result)

    confirm = input("💾 Do you want to save the merged result? (y/n): ").strip().lower()
    if confirm == 'y':
        save_path = os.path.join(DOWNLOADS_PATH, os.path.basename(file1).rsplit('.', 1)[0] + "_merged.txt")
        with open(save_path, 'w', encoding='utf-8') as f:
            f.write(merged)
        print(f"✅ Merged file saved at: {save_path}")
    else:
        print("❌ Merge was canceled. Nothing saved.")

    return True

def merge_docx_files(file1, file2, output_path):
    doc1 = Document(file1)
    doc2 = Document(file2)

    lines1 = [para.text.strip() for para in doc1.paragraphs if para.text.strip()]
    lines2 = [para.text.strip() for para in doc2.paragraphs if para.text.strip()]

    if lines1 == lines2:
        print("🔁 DOCX files are identical. No merge needed.")
        ask_file_deletion(file1, file2)
        return False

    print("🔍 DOCX files differ. Showing character-level changes:")
    combined = difflib.unified_diff(lines1, lines2, lineterm="")
    for line in combined:
        print(line)

    print("✏️ Merging line-by-line.")
    merged_doc = Document()
    for i in range(max(len(lines1), len(lines2))):
        a = lines1[i] if i < len(lines1) else ""
        b = lines2[i] if i < len(lines2) else ""
        if a == b:
            merged_doc.add_paragraph(a)
        else:
            print(f"Line {i+1} - A: {a}\n              B: {b}")
            choice = input("Choose (1) A, (2) B, (3) both, (4) custom: ").strip()
            if choice == '1':
                merged_doc.add_paragraph(a)
            elif choice == '2':
                merged_doc.add_paragraph(b)
            elif choice == '3':
                merged_doc.add_paragraph(a)
                merged_doc.add_paragraph(b)
            elif choice == '4':
                merged_doc.add_paragraph(input("Enter custom: "))

    confirm = input("💾 Save merged DOCX? (y/n): ").strip().lower()
    if confirm == 'y':
        merged_doc.save(output_path)
        print(f"✅ Saved: {output_path}")
    else:
        print("❌ Merge skipped.")
    return True

def merge_pdf_files(file1, file2, output_path):
    try:
        reader1 = PyPDF2.PdfReader(file1)
        reader2 = PyPDF2.PdfReader(file2)
        writer = PyPDF2.PdfWriter()

        text1 = "\n".join(page.extract_text() or "" for page in reader1.pages)
        text2 = "\n".join(page.extract_text() or "" for page in reader2.pages)

        if text1 == text2:
            print("🔁 PDF files are identical. No merge needed.")
            ask_file_deletion(file1, file2)
            return False

        print("📄 PDF files differ. Merging by unique content.")
        seen = set()
        combined_text = []
        for t in (text1.splitlines() + text2.splitlines()):
            if t not in seen:
                seen.add(t)
                combined_text.append(t)

        packet = BytesIO()
        can = canvas.Canvas(packet, pagesize=letter)
        y = 750
        for line in combined_text:
            if y < 50:
                can.showPage()
                y = 750
            can.drawString(50, y, line[:100])
            y -= 15
        can.save()
        packet.seek(0)
        new_pdf = PyPDF2.PdfReader(packet)
        writer.append_pages_from_reader(new_pdf)

        confirm = input("💾 Save merged PDF? (y/n): ").strip().lower()
        if confirm == 'y':
            with open(output_path, "wb") as out:
                writer.write(out)
            print(f"✅ PDF saved: {output_path}")
        else:
            print("❌ Merge canceled.")
    except Exception as e:
        print(f"⚠️ Error merging PDF: {e}")
    return True


def detect_and_handle_duplicates(target_folder):
    print("\n🔎 Scanning for duplicate files by type...\n")

    file_categories = {
        "text": ['.txt', '.md', '.py', '.csv'],
        "pdf": ['.pdf'],
        "word": ['.docx']
    }

    duplicates_by_type = {cat: [] for cat in file_categories}
    seen_hashes = {}

    for root, dirs, files in os.walk(target_folder):
        for file in files:
            file_path = os.path.join(root, file)
            ext = os.path.splitext(file)[1].lower()
            category = next((cat for cat, ext_list in file_categories.items() if ext in ext_list), None)
            if not category:
                continue

            file_hash = calculate_hash(file_path)
            if file_hash in seen_hashes:
                duplicates_by_type[category].append((seen_hashes[file_hash], file_path))
            else:
                seen_hashes[file_hash] = file_path

    print("📊 Duplicate Summary by Type:")
    for cat, dups in duplicates_by_type.items():
        print(f" - {cat.title()} Files: {len(dups)} duplicate pairs")

    while True:
        choice = input("\n❓ Which file type would you like to scan and merge? (Text / PDF / Word) or type 'no' to skip: ").strip().lower()
        if choice == "no":
            print("✅ Skipping remaining duplicate types.\n")
            break

        cat = {"text": "text", "pdf": "pdf", "word": "word"}.get(choice)
        if not cat or not duplicates_by_type[cat]:
            print("❌ Invalid or empty category.")
            continue

        duplicates = duplicates_by_type[cat]
        print(f"\n⚡ Found {len(duplicates)} duplicate {cat} file pairs.\n")

        set_size = 5
        auto_merged_count = 0

        for i, (file1, file2) in enumerate(duplicates, start=1):
            print(f"\n🛑 Duplicate Pair {i} of {len(duplicates)}:\n - {file1}\n - {file2}")
            ext = os.path.splitext(file1)[1].lower()
            base_name = os.path.basename(file1).rsplit('.', 1)[0]
            save_path = os.path.join(DOWNLOADS_PATH, f"{base_name}_merged{ext}")

            if ext == ".docx":
                user_decision_made = merge_docx_files(file1, file2, save_path)
            elif ext == ".pdf":
                user_decision_made = merge_pdf_files(file1, file2, save_path)
            else:
                user_decision_made = compare_and_merge(file1, file2)

            if not user_decision_made:
                auto_merged_count += 1

            if i % set_size == 0 and i != len(duplicates):
                if auto_merged_count == set_size:
                    print(f"\n✅ Last {set_size} files were identical and auto-merged.")
                else:
                    print(f"\n⏭️ You've reviewed {set_size} duplicate files.")
                    choice = input(f"❓ Continue with next {set_size}? (y/n): ").strip().lower()
                    if choice != 'y':
                        print("❌ Stopping duplicate merging early for this type.")
                        break
                auto_merged_count = 0

        print(f"✅ Finished processing {cat.title()} file duplicates.\n")
        cont = input("❓ Do you want to process another file type? (y/n): ").strip().lower()
        if cont != 'y':
            break


def psychological_quiz():
    questions = [
        "Do you feel anxious deleting old files?",
        "Do you often save things 'just in case' you might need them later?",
        "Do you have multiple copies of the same document or picture?",
        "Is your desktop cluttered with many files and shortcuts?",
        "Do you find it hard to organize your folders?",
        "Have you ever bought extra storage because you ran out of space?",
        "Do you rarely clean or sort your Downloads folder?",
        "Do you keep old versions of files you don't use anymore?",
        "Do you feel overwhelmed when trying to clean your computer?",
        "Do you avoid deleting things even if you know they're not important?"
    ]

    score = 0
    print("\n🧠 PSYCHOLOGICAL QUIZ")
    print("Answer each question: (0 = Never, 1 = Sometimes, 2 = Always)\n")

    for i, question in enumerate(questions, start=1):
        while True:
            try:
                ans = int(input(f"Q{i}: {question} [0-2]: "))
                if ans in [0, 1, 2]:
                    score += ans
                    break
                else:
                    print("❌ Please enter 0, 1, or 2 only.")
            except:
                print("❌ Invalid input. Please enter 0, 1, or 2.")
    return score

def analyze_final(system_score, psychological_score):
    print("\n📊 FINAL DIGITAL HOARDING RISK REPORT\n")

    if system_score <= 1:
        system_risk = "Normal"
    elif system_score == 2:
        system_risk = "Borderline"
    else:
        system_risk = "Severe"

    if psychological_score <= 5:
        psych_risk = "Low Emotional Hoarding"
    elif psychological_score <= 12:
        psych_risk = "Medium Emotional Hoarding"
    else:
        psych_risk = "High Emotional Hoarding"

    if system_risk == "Severe" or psych_risk == "High Emotional Hoarding":
        final_risk = "🚨 Severe Digital Hoarder"
    elif system_risk == "Borderline" or psych_risk == "Medium Emotional Hoarding":
        final_risk = "⚠️ Borderline Digital Hoarder"
    else:
        final_risk = "✅ Normal User"

    print(f"🖥️ System Behavior Risk: {system_risk}")
    print(f"🧠 Psychological Behavior Risk: {psych_risk}")
    print(f"\n🎯 FINAL RISK LEVEL: {final_risk}\n")

    if final_risk == "✅ Normal User":
        print("✅ You have healthy digital habits. Keep organizing your files regularly!")
    elif final_risk == "⚠️ Borderline Digital Hoarder":
        print("⚠️ You're starting to accumulate too many files.\n- Clean up old files.\n- Organize folders.\n- Delete unused apps.\n- Backup important data.")
    else:
        print("🚨 Severe Digital Hoarding Detected!\n- Immediately delete duplicates.\n- Uninstall unused applications.\n- Backup and organize data.\n- Use tools like CCleaner, Gemini, fdupes.\n- Consider a fresh system reinstall if performance is very poor.")

def run_all_merges_and_quiz():
    detect_and_handle_duplicates(DOWNLOADS_PATH)
    psych_score = psychological_quiz()
    analyze_final(system_points, psych_score)

# ======== MAIN PROGRAM ========

if __name__ == "__main__":
    print("🔎 SCANNING SYSTEM...\n")
    scan_normal_filesystem(DOWNLOADS_PATH)

    system_points = 0
    if len(NORMAL_FILES) > 50000:
        system_points += 1
    if len(NORMAL_FOLDERS) > 10000:
        system_points += 1
    if TOTAL_FILES_IN_ARCHIVES > 10000:
        system_points += 1
    if MAX_NESTING_DEPTH > 5:
        system_points += 1

    print("\n✅ SYSTEM SCAN COMPLETE!")
    print(f"Total Normal Files: {len(NORMAL_FILES)}")
    print(f"Total Normal Folders: {len(NORMAL_FOLDERS)}")
    print(f"Files Inside Archives: {TOTAL_FILES_IN_ARCHIVES}")
    print(f"Folders Inside Archives: {TOTAL_FOLDERS_IN_ARCHIVES}")
    print(f"Deepest Nesting Level: {MAX_NESTING_DEPTH}")

    detect_and_handle_duplicates(DOWNLOADS_PATH)
    psych_score = psychological_quiz()
    analyze_final(system_points, psych_score)
